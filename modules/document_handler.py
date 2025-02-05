import os
import json
import logging
import hashlib
from haystack.schema import Document

# Megengedett fájltípusok
ALLOWED_FILE_TYPES = {".pdf", ".docx", ".txt", ".json", ".jsonl"}

def load_hashes(hash_file):
    """Hash-ek betöltése fájlból."""
    if os.path.exists(hash_file):
        try:
            with open(hash_file, "r") as f:
                return set(json.load(f))
        except Exception as e:
            logging.error(f"Hiba a hash fájl betöltésekor: {e}")
    return set()

def save_hashes(hashes, hash_file):
    """Hash-ek mentése fájlba úgy, hogy a meglévő adatokat ne írjuk felül."""
    # Betöltjük a meglévő hash-eket, ha a fájl létezik
    existing_hashes = set()
    if os.path.exists(hash_file):
        try:
            with open(hash_file, "r") as f:
                existing_hashes = set(json.load(f))
        except Exception as e:
            logging.error(f"Hiba a meglévő hash-ek betöltésekor: {e}")
    
    # Egyesítjük a meglévő és új hash-eket
    all_hashes = existing_hashes.union(hashes)

    # Új hash-ek mentése
    try:
        with open(hash_file, "w") as f:
            json.dump(list(all_hashes), f)
        logging.info(f"Hash-ek sikeresen mentve a fájlba: {hash_file}")
    except Exception as e:
        logging.error(f"Hiba a hash fájl mentésekor: {e}")

def calculate_hash(content):
    """Egyedi hash kiszámítása a dokumentum tartalmáról."""
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

def read_file(file_path):
    """Fájl beolvasása és tartalom visszaadása."""
    _, ext = os.path.splitext(file_path)
    try:
        if ext not in ALLOWED_FILE_TYPES:
            raise ValueError(f"Nem támogatott fájltípus: {ext}")

        if ext == ".pdf":
            import fitz
            with fitz.open(file_path) as pdf:
                return "\n".join([page.get_text() for page in pdf])
        elif ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        elif ext in {".json", ".jsonl"}:
            with open(file_path, "r", encoding="utf-8") as file:
                data = json.load(file) if ext == ".json" else [json.loads(line) for line in file]
                if isinstance(data, list):
                    return "\n\n".join([f"{item.get('title', '')}\n{item.get('text', '')}" for item in data])
                elif isinstance(data, dict):
                    return f"{data.get('title', '')}\n{data.get('text', '')}"
        else:
            raise ValueError(f"Nem támogatott fájltípus: {ext}")
    except Exception as e:
        logging.error(f"Hiba a fájl feldolgozásakor: {file_path} - {e}")
        return None

def index_documents_recursive(path, document_store, document_hashes, retriever, batch_size, hash_file):
    """Dokumentumok rekurzív indexelése, csak az új fájlokat hozzáadva."""
    files = []
    for root, _, file_names in os.walk(path):
        for file_name in file_names:
            file_path = os.path.join(root, file_name)
            _, ext = os.path.splitext(file_name)
            if ext in ALLOWED_FILE_TYPES:
                files.append(file_path)

    logging.info(f"{len(files)} fájl található a mappában. Ellenőrzés kezdése...")
    new_documents = []

    for file_path in files:
        content = read_file(file_path)
        if not content:
            logging.warning(f"Hibás vagy üres fájl kihagyva: {file_path}")
            continue

        doc_hash = calculate_hash(content)
        if doc_hash in document_hashes:
            logging.info(f"Dokumentum már indexelve: {file_path}")
            continue
        document_hashes.add(doc_hash)

        # Új dokumentum feldolgozása
        document_hashes.add(doc_hash)
        document = Document(content=content, meta={"name": os.path.basename(file_path)})
        new_documents.append(document)

    if new_documents:
        logging.info(f"{len(new_documents)} új dokumentum található, indexelés batch mérete: {batch_size}")

        # Csoportos indexelés
        for i in range(0, len(new_documents), batch_size):
            batch = new_documents[i:i + batch_size]
            logging.info(f"⚡ {len(batch)} dokumentum feldolgozása ({i + 1}-{i + len(batch)})...")
            document_store.write_documents(batch)
            document_store.update_embeddings(retriever, batch_size=batch_size)

        # Hash-ek mentése
        save_hashes(document_hashes, hash_file)
        logging.info("✅ Az összes új dokumentum sikeresen hozzáadva az adatbázishoz.")
    else:
        logging.info("Nincsenek új dokumentumok.")

